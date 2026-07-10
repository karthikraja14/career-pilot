#!/usr/bin/env python3
"""
DOCX Resume Export — Generate a Word document version of your resume.

Some company portals prefer .docx over PDF. This generates an ATS-friendly
single-column Word document from the same resume data.

Usage:
  python docx_export.py
  python docx_export.py --output output/My_Resume.docx
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


def build_docx_resume(resume_data: dict, output_path: str) -> str:
    """Build an ATS-friendly DOCX resume."""
    if Document is None:
        print("  Install python-docx:  pip install python-docx")
        return ""

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_data["name"])
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(25, 60, 82)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(resume_data.get("title", ""))
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(100, 100, 100)

    # Contact
    contact = resume_data.get("contact", {})
    contact_parts = [contact.get(k, "") for k in ["phone", "email", "linkedin", "location"] if contact.get(k)]
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run("  |  ".join(contact_parts))
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(100, 100, 100)

    # Summary
    _add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary_para = doc.add_paragraph(resume_data.get("summary", ""))
    summary_para.style.font.size = Pt(10)

    # Skills
    _add_section_heading(doc, "TECHNICAL SKILLS")
    for category, skills_text in resume_data.get("skills", {}).items():
        p = doc.add_paragraph()
        cat_run = p.add_run(f"{category}: ")
        cat_run.bold = True
        cat_run.font.size = Pt(10)
        cat_run.font.color.rgb = RGBColor(25, 60, 82)
        skills_run = p.add_run(skills_text)
        skills_run.font.size = Pt(10)

    # Experience
    _add_section_heading(doc, "WORK EXPERIENCE")
    for job in resume_data.get("experience", []):
        # Role + Period
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(job["role"])
        role_run.bold = True
        role_run.font.size = Pt(11)
        role_run.font.color.rgb = RGBColor(25, 60, 82)
        role_para.add_run(f"  —  {job['period']}").font.size = Pt(9)

        # Company
        company_para = doc.add_paragraph(job["company"])
        company_para.style.font.size = Pt(10)

        # Bullets
        for bullet in job.get("bullets", []):
            bp = doc.add_paragraph(bullet, style="List Bullet")
            bp.style.font.size = Pt(10)

    # Achievements
    if resume_data.get("achievements"):
        _add_section_heading(doc, "KEY ACHIEVEMENTS")
        for ach in resume_data["achievements"]:
            ap = doc.add_paragraph(ach, style="List Bullet")
            ap.style.font.size = Pt(10)

    # Education
    _add_section_heading(doc, "EDUCATION")
    for edu in resume_data.get("education", []):
        ep = doc.add_paragraph()
        deg_run = ep.add_run(edu["degree"])
        deg_run.bold = True
        deg_run.font.size = Pt(10)
        ep.add_run(f"  —  {edu['institution']}  ({edu['year']})").font.size = Pt(9)

    # Languages
    if resume_data.get("languages"):
        _add_section_heading(doc, "LANGUAGES")
        doc.add_paragraph("  |  ".join(resume_data["languages"]))

    # Save
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def _add_section_heading(doc, title: str):
    """Add a section heading with underline."""
    para = doc.add_paragraph()
    para.space_before = Pt(12)
    para.space_after = Pt(4)
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(25, 60, 82)
    # Add a thin border below (via paragraph border)
    from docx.oxml.ns import qn
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '193C52',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def run_docx_export(output_path: str = "output/Resume.docx"):
    """Main entry point."""
    try:
        from resume_data import RESUME_DATA
    except ImportError:
        print("  Error: resume_data.py not found.")
        return

    if Document is None:
        print("  Install python-docx:  pip install python-docx")
        return

    path = build_docx_resume(RESUME_DATA, output_path)
    if path:
        size_kb = os.path.getsize(path) / 1024
        print(f"\n  DOCX resume saved: {path} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Export resume as DOCX")
    parser.add_argument("--output", default="output/Resume.docx")
    args = parser.parse_args()
    run_docx_export(args.output)
